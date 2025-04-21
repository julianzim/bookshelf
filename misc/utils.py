import statistics
import logging
import os
import pytz

from docx import Document
from datetime import datetime

from src.reviews.models import Reviews



async def localize_reviews(reviews: list[Reviews], timezone: pytz.timezone):
    localized_reviews = []
    for review in reviews:
        created_at = review.created_at
        
        if created_at.tzinfo is None:
            created_at = pytz.UTC.localize(created_at)

        localized_created_at = created_at.astimezone(timezone)

        updated_review = review._asdict()
        updated_review['created_at'] = localized_created_at

        localized_reviews.append(updated_review)
    
    return localized_reviews


async def get_reviews_statistics(reviews: list[Reviews]):
    ratings = [review.rating for review in reviews]
    reviews_texts = [review.text for review in reviews]
    if ratings:
        average_rating = round(statistics.mean(ratings), 1)
    else:
        average_rating = 0
    ratings_count = len(ratings)
    reviews_count = len(reviews_texts)
        
    return {
        "average_rating": average_rating,
        "ratings_count": ratings_count,
        "reviews_count": reviews_count
    }


def get_logger(name: str = None, log_level: str = None, set_sqla_logger: bool = False):
    """
    Returns a configured logger with the specified name and logging level.
    The logging level can be set either via an argument or the LOG_LEVEL environment variable.
    """
    logger = logging.getLogger(name)
    if not logger.hasHandlers():
        log_format = "%(asctime)s - %(levelname)-8s - (%(name)s).%(funcName)s(%(lineno)d) - %(message)s"
        formatter = logging.Formatter(log_format)

        file_handler = logging.FileHandler("app.log")
        file_handler.setFormatter(formatter)

        console_handler = logging.StreamHandler()
        console_handler.setFormatter(formatter)

        if log_level is None:
            log_level = os.getenv("LOG_LEVEL", "INFO").upper()

        if isinstance(log_level, str):
            log_level = getattr(logging, log_level.upper(), logging.INFO)
        elif isinstance(log_level, int):
            pass
        else:
            raise ValueError("Log level must be either a string or an integer.")
        
        logger.setLevel(log_level)
        logger.addHandler(file_handler)
        logger.addHandler(console_handler)

        if set_sqla_logger:
            sqlalchemy_logger = logging.getLogger("sqlalchemy.engine")
            sqlalchemy_logger.setLevel(logging.WARNING)
            sqlalchemy_logger.addHandler(file_handler)

    return logger


def format_runs(runs):
    formatted_text = ""
    
    for run in runs:
        text = run.text
        if run.bold:
            text = f'<b class="bold-text">{text}</b>'
        if run.italic:
            text = f'<i class="italic-text">{text}</i>'
        if run.underline:
            text = f'<u class="underline-text">{text}</u>'
        formatted_text += text
    
    return formatted_text


def parse_readtime_pubdata(input_string: str):
    if not input_string:
        read_time, pub_date = 5, datetime.now()
    else:
        read_time_string, date_string = input_string.split(", ")
        try:
            read_time = int(read_time_string.split(": ")[1])
        except ValueError:
            read_time = 5
        try:
            pub_date = datetime.strptime(date_string.split(": ")[1], "%d.%m.%Y")
        except ValueError:
            pub_date = datetime.now()

    return {"read_time": read_time, "pub_date": pub_date}


def convert_docx_to_html(input_file: str):
    doc = Document(input_file)

    info_runs_str = "".join([run.text for run in doc.paragraphs[0].runs])
    result = parse_readtime_pubdata(info_runs_str)
    result["summary"] = "".join([run.text for run in doc.paragraphs[1].runs])
    result["html_content"] = ""
    result["theme"] = "theme"
    result["title"] = "title"
    level = 0

    for paragraph in doc.paragraphs[2:]:
        if not paragraph.text.strip():
            continue
        if paragraph.style.name.startswith("Heading"):
            level = int(paragraph.style.name[-1])
            if level == 1:
                result["theme"] = "".join([run.text for run in paragraph.runs])
            elif level == 2:
                result["title"] = "".join([run.text for run in paragraph.runs])
            else:
                result["html_content"] += ((level-1) * "\t" + f'<h{level} class="article-header{level}">{format_runs(paragraph.runs)}</h{level}>\n')
        else:
            if paragraph.runs != "":
                result["html_content"] += (level * "\t" + f'<p class="article-paragraph">{format_runs(paragraph.runs)}</p>\n')
            else:
                result["html_content"] += ("\n")
    
    return result
