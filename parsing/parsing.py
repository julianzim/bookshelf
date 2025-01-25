from docx import Document
from pathlib import Path
import os
from datetime import datetime
from pprint import pprint


def format_runs(runs):
    formatted_text = ""
    for run in runs:
        text = run.text
        if run.bold:
            text = f'<b class="bold-text">{text}</b>'
        if run.italic:
            text = f'<i class="italic-text">{text}</i>'
        if run.underline:
            text = f'<u class="underline-text">{text}</u>'
        formatted_text += text
    
    return formatted_text


def parse_readtime_pubdata(input_string):
    if not input_string:
        read_time, pub_date = 5, datetime.now()
    else:
        read_time_string, date_string = input_string.split(", ")
        try:
            read_time = int(read_time_string.split(": ")[1])
        except ValueError:
            read_time = 5
        try:
            pub_date = datetime.strptime(date_string.split(": ")[1], "%d.%m.%Y")
        except ValueError:
            pub_date = datetime.now()

    return {"read_time": read_time, "pub_date": pub_date}


def convert_docx_to_html(input_file: str):
    doc = Document(input_file)

    result = parse_readtime_pubdata(doc.paragraphs[0].runs[0].text)
    result["summary"] = doc.paragraphs[1].runs[0].text
    result["html_content"] = ""
    level = 0

    for paragraph in doc.paragraphs[2:]:
        if not paragraph.text.strip():
            continue

        if paragraph.style.name.startswith("Heading"):
            level = int(paragraph.style.name[-1])
            if level == 1:
                result["theme"] = paragraph.runs[0].text
            elif level == 2:
                result["title"] = paragraph.runs[0].text
            else:
                result["html_content"] += ((level-1) * "\t" + f'<h{level} class="article-header{level}">{format_runs(paragraph.runs)}</h{level}>\n')
        else:
            if paragraph.runs != "":
                result["html_content"] += (level * "\t" + f'<p class="article-paragraph">{format_runs(paragraph.runs)}</p>\n')
            else:
                result["html_content"] += ("\n")

    with open("parsing/article.html", "w", encoding="utf-8") as f:
        f.write(result["html_content"])
    
    pprint(result)


if __name__ == "__main__":
    input_path = os.path.abspath("parsing/article.docx")

    if Path(input_path).suffix.lower() != ".docx":
        print("Скрипт поддерживает только .docx файлы")
    else:
        convert_docx_to_html(input_path)
        print(f"Файл успешно преобразован в HTML и сохранён в \"parsing/article.html\"")
